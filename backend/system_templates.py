"""System-seeded DOCX templates for each industry/subdomain.

These are built once on first server start and made available read-only to
every user via /api/templates (with is_system=true flag). Users can also
upload their own templates as before.

Templates are generated in-memory using python-docx with Romanian gas
engineering boilerplate text + placeholders. Replace these later with
your real legal forms.
"""
import io
import base64
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _h(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)


def _p(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    if bold:
        r.bold = True


def _build_cerere_racordare() -> bytes:
    d = Document()
    _h(d, "CERERE DE RACORDARE LA REȚEAUA DE DISTRIBUȚIE GAZE NATURALE")
    d.add_paragraph()
    _p(d, "Către: {{osd}}")
    d.add_paragraph()
    _p(d, "Subsemnatul/Subscrisa {{beneficiar}}, cu domiciliul/sediul în localitatea {{localitate}}, județul {{judet}}, adresa {{adresa_lucrare}}, telefon {{telefon}}, email {{email}}, vă rog să aprobați racordarea la rețeaua de distribuție gaze naturale a imobilului situat la adresa menționată.")
    d.add_paragraph()
    _p(d, "Date tehnice ale lucrării:", bold=True)
    _p(d, "• Tip lucrare: {{tip_lucrare}}")
    _p(d, "• Debit instalat: {{debit_instalat}} mc/h")
    _p(d, "• Putere instalată: {{putere_instalata_kw}} kW")
    _p(d, "• Lungime estimată branșament: {{lungime_bransament}} m")
    _p(d, "• Presiune regim solicitată: {{presiune_regim}}")
    _p(d, "• Contor propus: {{contor_orientativ}}")
    d.add_paragraph()
    _p(d, "Date contract: nr. {{numar_contract}} din {{data_contract}}.")
    _p(d, "Proiectant: {{proiectant}}")
    _p(d, "Executant: {{executant}}")
    d.add_paragraph()
    _p(d, "Observații: {{observatii}}")
    d.add_paragraph()
    _p(d, "Data: {{data_document}}")
    _p(d, "Semnătura beneficiar: ________________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_memoriu_tehnic() -> bytes:
    d = Document()
    _h(d, "MEMORIU TEHNIC — INSTALAȚIE GAZE NATURALE", size=16)
    d.add_paragraph()
    _p(d, "1. DATE GENERALE", bold=True)
    _p(d, "Beneficiar: {{beneficiar}}")
    _p(d, "Adresa lucrării: {{adresa_lucrare}}, {{localitate}}, jud. {{judet}}")
    _p(d, "Telefon contact: {{telefon}} | Email: {{email}}")
    _p(d, "Operator distribuție (OSD): {{osd}}")
    _p(d, "Tip lucrare: {{tip_lucrare}}")
    _p(d, "Contract: nr. {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "2. ECHIPA TEHNICĂ AUTORIZATĂ", bold=True)
    _p(d, "Proiectant ANRE: {{proiectant}}")
    _p(d, "Executant ANRE: {{executant}}")
    _p(d, "Verificator documentație (VGD): {{verificator_vgd}}")
    _p(d, "Responsabil tehnic execuție (RTE): {{responsabil_rte}}")
    d.add_paragraph()
    _p(d, "3. DATE TEHNICE", bold=True)
    _p(d, "Debit instalat: {{debit_instalat}} mc/h")
    _p(d, "Debit calculat: {{debit_calculat_mc_h}} mc/h")
    _p(d, "Debit recomandat (cu marjă 10%): {{debit_recomandat_mc_h}} mc/h")
    _p(d, "Putere instalată estimată: {{putere_instalata_kw}} kW")
    _p(d, "Presiune regim: {{presiune_regim}}")
    _p(d, "Diametru conductă: {{diametru_conducta}}")
    _p(d, "Material conductă: {{material_conducta}}")
    _p(d, "Lungime branșament: {{lungime_bransament}} m")
    _p(d, "Punct racordare: {{punct_racordare}}")
    _p(d, "Post reglare: {{post_reglare}}")
    _p(d, "Contor recomandat: {{contor_orientativ}}")
    _p(d, "Risc presiune: {{risc_presiune}}")
    _p(d, "Estimare cost branșament: {{estimare_cost}} RON")
    d.add_paragraph()
    _p(d, "4. OBSERVAȚII TEHNICE", bold=True)
    _p(d, "{{observatii}}")
    d.add_paragraph()
    _p(d, "Data emiterii: {{data_document}}")
    _p(d, "Întocmit: {{proiectant}}")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_borderou() -> bytes:
    d = Document()
    _h(d, "BORDEROU DOCUMENTE — DOSAR TEHNIC")
    d.add_paragraph()
    _p(d, "Beneficiar: {{beneficiar}}")
    _p(d, "Adresa lucrare: {{adresa_lucrare}}, {{localitate}}, {{judet}}")
    _p(d, "Tip lucrare: {{tip_lucrare}}")
    _p(d, "Contract: {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "Conținut dosar:", bold=True)
    _p(d, "1. Cerere de racordare")
    _p(d, "2. Memoriu tehnic")
    _p(d, "3. Plan de încadrare (anexă)")
    _p(d, "4. Plan de situație (anexă)")
    _p(d, "5. Schemă izometrică (anexă)")
    _p(d, "6. Acord acces / declarație proprietate (anexă)")
    _p(d, "7. Verificare VGD: {{verificator_vgd}}")
    _p(d, "8. Confirmare RTE: {{responsabil_rte}}")
    d.add_paragraph()
    _p(d, "Întocmit: {{proiectant}} | Data: {{data_document}}")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_adresa_osd() -> bytes:
    d = Document()
    _h(d, "ADRESĂ CĂTRE OPERATORUL SISTEMULUI DE DISTRIBUȚIE", size=14)
    d.add_paragraph()
    _p(d, "Către: {{osd}}")
    _p(d, "În atenția: Departamentul Racordări")
    d.add_paragraph()
    _p(d, "Stimată/Stimate domnule director,")
    d.add_paragraph()
    _p(d, "Prin prezenta, vă transmitem documentația tehnică aferentă lucrării de {{tip_lucrare}} pentru beneficiarul {{beneficiar}}, situat la adresa {{adresa_lucrare}}, {{localitate}}, județul {{judet}}.")
    d.add_paragraph()
    _p(d, "Detalii lucrare:", bold=True)
    _p(d, "• Contract nr. {{numar_contract}} din {{data_contract}}")
    _p(d, "• Debit instalat: {{debit_instalat}} mc/h")
    _p(d, "• Lungime branșament: {{lungime_bransament}} m")
    _p(d, "• Proiectant ANRE: {{proiectant}}")
    _p(d, "• Executant ANRE: {{executant}}")
    d.add_paragraph()
    _p(d, "Vă rugăm să analizați documentația și să comunicați avizul dvs. în termenul legal.")
    d.add_paragraph()
    _p(d, "Cu deosebită stimă,")
    _p(d, "{{proiectant}}")
    _p(d, "Data: {{data_document}}")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_certificare_vgd() -> bytes:
    d = Document()
    _h(d, "CERTIFICARE INTERNĂ — VERIFICATOR DOCUMENTAȚIE (VGD)", size=14)
    d.add_paragraph()
    _p(d, "Beneficiar lucrare: {{beneficiar}}")
    _p(d, "Adresa lucrării: {{adresa_lucrare}}, {{localitate}}, {{judet}}")
    _p(d, "Tip lucrare: {{tip_lucrare}}")
    _p(d, "Contract: nr. {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "1. VERIFICATOR DOCUMENTAȚIE", bold=True)
    _p(d, "Nume: {{verificator_vgd}}")
    _p(d, "Atestat ANRE: {{atestat_vgd}}")
    _p(d, "Data verificării: {{data_verificare_vgd}}")
    _p(d, "Status verificare: {{status_vgd}}")
    d.add_paragraph()
    _p(d, "2. OBSERVAȚII VERIFICARE", bold=True)
    _p(d, "{{observatii_vgd}}")
    d.add_paragraph()
    _p(d, "Subsemnatul, în calitate de Verificator Documentație autorizat ANRE, certific verificarea completă a documentației tehnice aferente lucrării de mai sus, conform reglementărilor în vigoare.")
    d.add_paragraph()
    _p(d, "Data certificării: {{data_document}}")
    _p(d, "Semnătură VGD: ________________________")
    _p(d, "Ștampilă: ________________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_certificare_rte() -> bytes:
    d = Document()
    _h(d, "CERTIFICARE INTERNĂ — RESPONSABIL TEHNIC EXECUȚIE (RTE)", size=14)
    d.add_paragraph()
    _p(d, "Beneficiar lucrare: {{beneficiar}}")
    _p(d, "Adresa lucrării: {{adresa_lucrare}}, {{localitate}}, {{judet}}")
    _p(d, "Tip lucrare: {{tip_lucrare}}")
    _p(d, "Executant: {{executant}}")
    _p(d, "Contract: nr. {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "1. RESPONSABIL TEHNIC EXECUȚIE", bold=True)
    _p(d, "Nume: {{responsabil_rte}}")
    _p(d, "Autorizație ANRE: {{autorizatie_rte}}")
    _p(d, "Data verificării execuției: {{data_verificare_rte}}")
    _p(d, "Status execuție: {{status_rte}}")
    d.add_paragraph()
    _p(d, "2. OBSERVAȚII EXECUȚIE", bold=True)
    _p(d, "{{observatii_rte}}")
    d.add_paragraph()
    _p(d, "Subsemnatul, în calitate de Responsabil Tehnic cu Execuția autorizat ANRE, certific execuția conformă a lucrării de mai sus și respectarea documentației tehnice verificate.")
    d.add_paragraph()
    _p(d, "Data certificării: {{data_document}}")
    _p(d, "Semnătură RTE: ________________________")
    _p(d, "Ștampilă: ________________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_cerere_racordare_fv() -> bytes:
    d = Document()
    _h(d, "CERERE DE RACORDARE LA REȚEAUA ELECTRICĂ DE DISTRIBUȚIE — INSTALAȚIE FOTOVOLTAICĂ", size=14)
    d.add_paragraph()
    _p(d, "Către: {{operator_distributie}}")
    _p(d, "În atenția: Departamentul Racordări Prosumatori / Producători")
    d.add_paragraph()
    _p(d, "Subsemnatul/Subscrisa {{beneficiar}}, cu domiciliul/sediul în {{localitate}}, jud. {{judet}}, "
          "adresa {{adresa_lucrare}}, telefon {{telefon}}, email {{email}}, "
          "solicit racordarea la rețeaua electrică de distribuție a unei instalații fotovoltaice "
          "cu putere instalată de {{fv_p_kwp}} kWp, încadrată în categoria ANRE {{fv_categorie_anre}} "
          "({{fv_categorie_label}}).")
    d.add_paragraph()
    _p(d, "1. DATE TEHNICE INSTALAȚIE FV", bold=True)
    _p(d, "• Putere instalată: {{fv_p_kwp}} kWp")
    _p(d, "• Număr panouri: {{fv_n_panouri}} × {{fv_p_panou_wp}} Wp")
    _p(d, "• Configurație: {{fv_n_string}} string-uri × {{fv_n_serie}} panouri/string")
    _p(d, "• Tensiune DC string (STC): {{fv_u_string_v}} V")
    _p(d, "• Invertor recomandat: {{fv_invertor_kw}} kW (raport DC/AC = {{fv_raport_dc_ac}})")
    _p(d, "• Producție anuală estimată: {{fv_productie_anuala_kwh}} kWh "
          "({{fv_productie_specifica}} kWh/kWp · iradiație {{fv_iradiatie_kwh_m2_an}} kWh/m²/an)")
    d.add_paragraph()
    _p(d, "2. REGIM DE FUNCȚIONARE SOLICITAT", bold=True)
    _p(d, "Regim: {{fv_regim}}")
    _p(d, "Tip racord: {{fv_tip_racord}}")
    _p(d, "Aviz necesar (conform ANRE Ord. 34/2024): {{fv_aviz_necesar}}")
    d.add_paragraph()
    _p(d, "{IF fv_p_kwp <= 10.8: Solicit înregistrarea ca prosumator casnic, cu compensare cantitativă a energiei livrate, conform Legii 220/2008 și Ord. ANRE 15/2022 actualizat. ELSE Solicit emiterea ATR conform procedurii standard pentru categoria ANRE corespunzătoare puterii instalate.}")
    d.add_paragraph()
    _p(d, "{IF fv_p_kwp > 200: ATENȚIE — instalație de tip parc fotovoltaic. Anexez studiu de soluție, acord mediu și solicitare licență producere energie electrică ANRE. ELSE Anexez documentația tehnică standard conform Ord. 34/2024.}")
    d.add_paragraph()
    _p(d, "3. PROIECTANT / EXECUTANT", bold=True)
    _p(d, "Proiectant atestat ANRE: {{proiectant}} ({{atestat_proiectant}})")
    _p(d, "Executant atestat ANRE: {{executant}} ({{atestat_executant}})")
    d.add_paragraph()
    _p(d, "Contract intern: nr. {{numar_contract}} / {{data_contract}}")
    _p(d, "Observații: {{observatii}}")
    d.add_paragraph()
    _p(d, "Data: {{data_document}}")
    _p(d, "Semnătură beneficiar: ________________________")
    _p(d, "Ștampilă proiectant: ________________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_memoriu_tehnic_fv() -> bytes:
    d = Document()
    _h(d, "MEMORIU TEHNIC — INSTALAȚIE FOTOVOLTAICĂ", size=16)
    d.add_paragraph()
    _p(d, "1. DATE GENERALE", bold=True)
    _p(d, "Beneficiar: {{beneficiar}}")
    _p(d, "Adresa lucrării: {{adresa_lucrare}}, {{localitate}}, jud. {{judet}}")
    _p(d, "Telefon / Email: {{telefon}} / {{email}}")
    _p(d, "Operator distribuție: {{operator_distributie}}")
    _p(d, "Contract: nr. {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "2. ÎNCADRARE ANRE", bold=True)
    _p(d, "Putere instalată: {{fv_p_kwp}} kWp")
    _p(d, "Categorie: {{fv_categorie_anre}} — {{fv_categorie_label}}")
    _p(d, "Regim funcționare: {{fv_regim}}")
    _p(d, "Tip racord: {{fv_tip_racord}}")
    _p(d, "Aviz/avize necesare: {{fv_aviz_necesar}}")
    d.add_paragraph()
    _p(d, "{IF fv_categorie_anre == C1: Instalație de tip PROSUMATOR CASNIC. Procedură simplificată conform ANRE Ord. 34/2024 — ATR + CR emise în max. 30 de zile. Beneficiarul are dreptul la compensare cantitativă a energiei livrate în rețea, conform Legii 220/2008 cu modificările ulterioare. ELSE Aplicabile prevederi extinse ANRE — vezi paragrafele de mai jos.}")
    d.add_paragraph()
    _p(d, "{IF fv_categorie_anre == C2: Instalație de tip PROSUMATOR NON-CASNIC. ATR cu studiu de soluție simplificat. Racord trifazat 3x400V obligatoriu. Compensare cantitativă disponibilă pentru excedent. ELSE  }")
    _p(d, "{IF fv_categorie_anre == C3: Instalație de tip PRODUCĂTOR MIC (27-200 kWp). Studiu de soluție complet, ATR cu cerințe extinse. Pentru P > 100 kWp este obligatorie obținerea licenței de producere energie electrică de la ANRE. Releu de protecție conform EN 50549-1 obligatoriu. ELSE  }")
    _p(d, "{IF fv_categorie_anre == C4: Instalație de tip PARC FOTOVOLTAIC (> 200 kWp). Racord pe MT 20 kV. Necesită: studiu de soluție complet, acord de mediu, licență producere ANRE, celulă MT 24 kV, telegestiune cu protocol Modbus/IEC 61850. ELSE  }")
    d.add_paragraph()
    _p(d, "3. CONFIGURAȚIE TEHNICĂ", bold=True)
    _p(d, "Panouri fotovoltaice: {{fv_n_panouri}} buc × {{fv_p_panou_wp}} Wp "
          "(module monocristaline TOPCon/HJT, eficiență ≥ 21.5%)")
    _p(d, "Configurație string-uri: {{fv_n_string}} string-uri × {{fv_n_serie}} module în serie")
    _p(d, "Tensiune string STC (25°C): {{fv_u_string_v}} V")
    _p(d, "Invertor: {{fv_invertor_kw}} kW (raport DC/AC = {{fv_raport_dc_ac}})")
    _p(d, "{IF fv_p_kwp <= 5: Conectare monofazată 230V admisă (P ≤ 5 kWp). ELSE Conectare trifazată 3x400V obligatorie.}")
    d.add_paragraph()
    _p(d, "4. CABLURI ȘI PROTECȚII", bold=True)
    _p(d, "Cablu DC (string → invertor): {{fv_cablu_dc_mm2}} mm² — tip H1Z2Z2-K solar 1500V "
          "(cădere de tensiune calculată: {{fv_cablu_dc_caderea_pct}}%, sub limita admisă 1%)")
    _p(d, "Cablu AC (invertor → tablou racordare): {{fv_cablu_ac_mm2}} mm² — tip CYY-F/N2XH 0.6/1 kV "
          "(curent calculat: {{fv_cablu_ac_curent_a}} A, cădere admisă 1.5% conf. NTI-TEL-007)")
    d.add_paragraph()
    _p(d, "Protecții electrice obligatorii (conform I7-2011 și SR EN IEC 62548):")
    _p(d, "{{fv_protectii_lista}}")
    d.add_paragraph()
    _p(d, "5. PRODUCȚIE ESTIMATĂ", bold=True)
    _p(d, "Iradiație zonă: {{fv_iradiatie_kwh_m2_an}} kWh/m²/an (sursa: PVGIS-SARAH3, JRC)")
    _p(d, "Performance Ratio (PR): 0.78 (mediu, include pierderi cabluri/invertor/mismatch/T°/murdărire)")
    _p(d, "Producție anuală estimată: {{fv_productie_anuala_kwh}} kWh")
    _p(d, "Producție specifică: {{fv_productie_specifica}} kWh/kWp")
    _p(d, "Factor de utilizare anual: {{fv_factor_utilizare_pct}}%")
    d.add_paragraph()
    _p(d, "{IF fv_p_kwp >= 100: ATENȚIE: instalație ≥ 100 kWp — sunt obligatorii smart-meter bidirecțional clasa 1, telegestiune Modbus/M-Bus și relee de protecție EN 50549-1 (anti-islanding, U/f, ROCOF). ELSE Echipamentele standard de protecție specificate la pct. 4 sunt suficiente.}")
    d.add_paragraph()
    _p(d, "6. ECHIPA TEHNICĂ AUTORIZATĂ", bold=True)
    _p(d, "Proiectant ANRE: {{proiectant}} ({{atestat_proiectant}})")
    _p(d, "Executant ANRE: {{executant}} ({{atestat_executant}})")
    _p(d, "Verificator documentație: {{verificator_vgd}}")
    _p(d, "Responsabil tehnic execuție (RTE): {{responsabil_rte}}")
    d.add_paragraph()
    _p(d, "7. OBSERVAȚII", bold=True)
    _p(d, "{{observatii}}")
    d.add_paragraph()
    _p(d, "Data emiterii: {{data_document}}")
    _p(d, "Întocmit: {{proiectant}}")
    _p(d, "Semnătură și ștampilă: ________________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_adresa_od_fv() -> bytes:
    d = Document()
    _h(d, "ADRESĂ CĂTRE OPERATORUL DE DISTRIBUȚIE — INSTALAȚIE FOTOVOLTAICĂ", size=13)
    d.add_paragraph()
    _p(d, "Către: {{operator_distributie}}")
    _p(d, "În atenția: Departamentul Racordări Prosumatori / Producători")
    d.add_paragraph()
    _p(d, "Stimată/Stimate domnule director,")
    d.add_paragraph()
    _p(d, "Prin prezenta, depunem documentația tehnică completă aferentă instalației fotovoltaice "
          "ce urmează a fi racordată pentru beneficiarul {{beneficiar}}, situat la adresa "
          "{{adresa_lucrare}}, {{localitate}}, jud. {{judet}}.")
    d.add_paragraph()
    _p(d, "Date sintetice:", bold=True)
    _p(d, "• Putere instalată: {{fv_p_kwp}} kWp — categorie ANRE {{fv_categorie_anre}}")
    _p(d, "• Regim funcționare: {{fv_regim}}")
    _p(d, "• Producție anuală estimată: {{fv_productie_anuala_kwh}} kWh")
    _p(d, "• Invertor recomandat: {{fv_invertor_kw}} kW")
    _p(d, "• Contract: nr. {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "{IF fv_p_kwp <= 10.8: Solicit emiterea simplificată a ATR și CR (procedură prosumator casnic ≤ 10.8 kWp), conform ANRE Ord. 34/2024 — termen legal 30 de zile. ELSE Solicit deschiderea procedurii standard de racordare cu studiu de soluție conform categoriei ANRE indicate.}")
    d.add_paragraph()
    _p(d, "{IF fv_p_kwp > 100: Anexez în mod suplimentar: studiu de soluție complet, schema unifilară aprobată VGD, fișa tehnică invertor și panouri, certificate de conformitate EN 50549-1, dovada inițiere licență producere ANRE. ELSE Anexez documentația standard: cerere racordare, memoriu tehnic, schema unifilară, fișe tehnice echipamente, declarație proprietate.}")
    d.add_paragraph()
    _p(d, "Echipa autorizată:")
    _p(d, "• Proiectant ANRE: {{proiectant}}")
    _p(d, "• Executant ANRE: {{executant}}")
    _p(d, "• Verificator documentație: {{verificator_vgd}}")
    d.add_paragraph()
    _p(d, "Vă mulțumim pentru disponibilitate și rămânem deschiși oricăror clarificări tehnice.")
    d.add_paragraph()
    _p(d, "Cu stimă,")
    _p(d, "{{proiectant}}")
    _p(d, "Data: {{data_document}}")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_dtac_gaz() -> bytes:
    """DTAC — Documentație Tehnică pentru obținerea Autorizației de Construire (gaze naturale).
    Conform Legii 50/1991 + Norma 84-2018 ANRE + HG 525/1996 PUG.
    """
    d = Document()
    _h(d, "DOCUMENTAȚIE TEHNICĂ PENTRU AUTORIZAȚIA DE CONSTRUIRE (D.T.A.C.)", size=14)
    _p(d, "INSTALAȚIE GAZE NATURALE — Legea 50/1991 + Ord. ANRE 89/2018", bold=True)
    d.add_paragraph()
    _p(d, "I. FOAIE DE CAPĂT", bold=True)
    _p(d, "Denumire investiție: {{denumire_investitie}}")
    _p(d, "Amplasament: {{adresa_lucrare}}, {{localitate}}, jud. {{judet}}")
    _p(d, "CF / Nr. Cadastral: {{numar_cadastral}} / {{numar_carte_funciara}}")
    _p(d, "Beneficiar: {{beneficiar}}")
    _p(d, "CUI/CNP: {{beneficiar_cui}}")
    _p(d, "Proiectant general: {{proiectant_general}}")
    _p(d, "Proiectant specialitate gaze: {{proiectant}} — Atestat ANRE {{atestat_proiectant}}")
    _p(d, "Faza: D.T.A.C. — Documentație Tehnică pentru Autorizație de Construire")
    _p(d, "Nr. proiect: {{numar_proiect}} / {{data_document}}")
    d.add_paragraph()
    _p(d, "II. BORDEROU PIESE SCRISE", bold=True)
    _p(d, "1. Foaie de capăt")
    _p(d, "2. Borderou piese scrise și desenate")
    _p(d, "3. Memoriu tehnic general")
    _p(d, "4. Memoriu tehnic specialitate gaze naturale")
    _p(d, "5. Breviar de calcul (debite, presiuni, dimensionare)")
    _p(d, "6. Caiet de sarcini și liste cu cantități")
    _p(d, "7. Devize generale și pe categorii de lucrări")
    _p(d, "8. Program control calitate execuție (PCCVI)")
    _p(d, "9. Referat verificare tehnică Vg (verificator atestat ANRE)")
    d.add_paragraph()
    _p(d, "III. BORDEROU PIESE DESENATE", bold=True)
    _p(d, "P1. Plan de încadrare în zonă (sc 1:5000 / 1:10000)")
    _p(d, "P2. Plan de situație (sc 1:500 / 1:1000)")
    _p(d, "P3. Plan parter / nivele cu traseu instalație gaze (sc 1:50 / 1:100)")
    _p(d, "P4. Schemă izometrică instalație (sc 1:50)")
    _p(d, "P5. Detalii branșament — secțiune transversală")
    _p(d, "P6. Detalii post reglare / contor")
    _p(d, "P7. Plan amplasament aparate consumatoare cu degajări de siguranță")
    d.add_paragraph()
    _p(d, "IV. MEMORIU TEHNIC GENERAL DTAC", bold=True)
    _p(d, "4.1 Descrierea generală a lucrării")
    _p(d, "Prezenta documentație are ca obiect lucrarea de {{tip_lucrare}} la imobilul situat la adresa {{adresa_lucrare}}, {{localitate}}, județul {{judet}}, în vederea racordării la rețeaua publică de distribuție gaze naturale a OSD-ului {{osd}}.")
    _p(d, "4.2 Încadrare urbanistică")
    _p(d, "Conform Certificatului de Urbanism nr. {{cu_numar}} / {{cu_data}} emis de Primăria {{primaria}}, zona de construcție permite execuția lucrărilor de instalații gaze naturale.")
    _p(d, "4.3 Soluția tehnică adoptată")
    _p(d, "• Debit total instalat: {{debit_instalat}} mc/h")
    _p(d, "• Putere termică totală: {{putere_instalata_kw}} kW")
    _p(d, "• Presiune regim: {{presiune_regim}}")
    _p(d, "• Tip rețea racord: {{punct_racordare}}")
    _p(d, "• Material conductă: {{material_conducta}}, dimensiune {{diametru_conducta}}")
    _p(d, "• Lungime branșament: {{lungime_bransament}} m")
    _p(d, "• Post reglare/măsurare: {{post_reglare}}")
    _p(d, "• Sistem măsurare: {{contor_orientativ}}")
    d.add_paragraph()
    _p(d, "V. RESPECTAREA NORMELOR DE SIGURANȚĂ", bold=True)
    _p(d, "Documentația respectă Ord. ANRE 89/2018 (Normă tehnică instalații gaze), Legea 123/2012, NTPEE-2018, Lg 50/1991, P118-99 (protecție incendiu) și HG 925/1995 (verificare proiecte).")
    d.add_paragraph()
    _p(d, "Întocmit: {{proiectant}}                     Verificat tehnic: {{verificator_vgd}} ({{atestat_vgd}})")
    _p(d, "Data: {{data_document}}                     Ștampilă proiectant: ________________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_dtoe_gaz() -> bytes:
    """DTOE — Documentație Tehnică de Organizare a Execuției lucrărilor."""
    d = Document()
    _h(d, "DOCUMENTAȚIE TEHNICĂ DE ORGANIZARE A EXECUȚIEI (D.T.O.E.)", size=14)
    _p(d, "Lucrări de instalații gaze naturale — execuție pe șantier", bold=True)
    d.add_paragraph()
    _p(d, "1. DATE GENERALE", bold=True)
    _p(d, "Beneficiar: {{beneficiar}}")
    _p(d, "Adresa: {{adresa_lucrare}}, {{localitate}}, jud. {{judet}}")
    _p(d, "Executant ANRE: {{executant}} — Atestat {{atestat_executant}}")
    _p(d, "RTE atestat ANRE: {{responsabil_rte}} ({{autorizatie_rte}})")
    _p(d, "Contract execuție: nr. {{numar_contract}} / {{data_contract}}")
    d.add_paragraph()
    _p(d, "2. ORGANIZAREA ȘANTIERULUI", bold=True)
    _p(d, "2.1 Suprafață ocupată temporar pentru organizare: {{suprafata_organizare}} mp")
    _p(d, "2.2 Acces utilaje și aprovizionare materiale: {{acces_santier}}")
    _p(d, "2.3 Surse de utilități pentru organizare: {{surse_utilitati}}")
    _p(d, "2.4 Spații de depozitare materiale: {{depozitare_materiale}}")
    _p(d, "2.5 Vestiare și grupuri sanitare: {{vestiare}}")
    d.add_paragraph()
    _p(d, "3. PROGRAM DE EXECUȚIE PE FAZE", bold=True)
    _p(d, "Faza 1 — Trasare și marcaj traseu conductă")
    _p(d, "Faza 2 — Săpături în profil normat (adâncime min. 0.9 m conform NTPEE-2018)")
    _p(d, "Faza 3 — Montaj conductă PE/OL și sudare/electrofuziune cu probe sudori certificați")
    _p(d, "Faza 4 — Probe de presiune (rezistență 1.5 × Pmax, etanșeitate 6h conform NTPEE-2018)")
    _p(d, "Faza 5 — Montaj robineți, post reglare, contor")
    _p(d, "Faza 6 — Probă tehnologică în prezența reprezentantului OSD")
    _p(d, "Faza 7 — Acoperire tranșee cu nisip + bandă avertizare + recompactare strat")
    _p(d, "Faza 8 — Refacere stare inițială sistematizare verticală")
    _p(d, "Durată totală estimată: {{durata_executie}} zile lucrătoare")
    d.add_paragraph()
    _p(d, "4. MĂSURI DE SECURITATE ȘI SĂNĂTATE", bold=True)
    _p(d, "Conform Lg 319/2006 SSM, HG 300/2006 șantiere, HG 1051/2006 echipamente, NTPEE-2018.")
    _p(d, "Echipa execuție va fi instruită zilnic. Operațiunile cu foc deschis necesită permis emis de RTE.")
    d.add_paragraph()
    _p(d, "5. MĂSURI DE PROTECȚIE A MEDIULUI", bold=True)
    _p(d, "Materialul rezultat din săpături — depozitat temporar și transportat la depozit autorizat. Refacerea stării inițiale a terenului este obligatorie.")
    d.add_paragraph()
    _p(d, "Întocmit: {{proiectant}}                       RTE: {{responsabil_rte}}")
    _p(d, "Data: {{data_document}}                       Semnătură & ștampilă: __________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_dispozitie_santier() -> bytes:
    """Dispoziție de Șantier — modificări față de proiectul aprobat în timpul execuției."""
    d = Document()
    _h(d, "DISPOZIȚIE DE ȘANTIER nr. {{ds_numar}} / {{data_document}}", size=14)
    _p(d, "Lucrare instalație gaze naturale", bold=True)
    d.add_paragraph()
    _p(d, "Lucrare: {{tip_lucrare}}")
    _p(d, "Beneficiar: {{beneficiar}}")
    _p(d, "Amplasament: {{adresa_lucrare}}, {{localitate}}, jud. {{judet}}")
    _p(d, "Contract: {{numar_contract}} / {{data_contract}}")
    _p(d, "Documentație inițială: nr. proiect {{numar_proiect}}, faza {{faza_proiect}}")
    d.add_paragraph()
    _p(d, "1. SITUAȚIA EXISTENTĂ ÎN PROIECT", bold=True)
    _p(d, "{{situatie_proiect_initial}}")
    d.add_paragraph()
    _p(d, "2. SITUAȚIA CONSTATATĂ PE TEREN", bold=True)
    _p(d, "{{situatie_constatata_teren}}")
    d.add_paragraph()
    _p(d, "3. MODIFICAREA PROPUSĂ", bold=True)
    _p(d, "{{modificare_propusa}}")
    d.add_paragraph()
    _p(d, "4. JUSTIFICAREA MODIFICĂRII", bold=True)
    _p(d, "{{justificare_modificare}}")
    d.add_paragraph()
    _p(d, "5. CONSECINȚE TEHNICE", bold=True)
    _p(d, "• Asupra siguranței: {{impact_siguranta}}")
    _p(d, "• Asupra costului lucrării: {{impact_cost}} RON (± față de devizul inițial)")
    _p(d, "• Asupra termenului: {{impact_termen}}")
    _p(d, "• Asupra autorizației de construire: {{impact_autorizatie}}")
    d.add_paragraph()
    _p(d, "6. SEMNĂTURI", bold=True)
    _p(d, "Proiectant: {{proiectant}}                         Semnătură: ____________________")
    _p(d, "Verificator tehnic Vg: {{verificator_vgd}}        Semnătură: ____________________")
    _p(d, "RTE: {{responsabil_rte}}                          Semnătură: ____________________")
    _p(d, "Beneficiar: {{beneficiar}}                        Semnătură: ____________________")
    _p(d, "Reprezentant OSD: ____________________             Semnătură: ____________________")
    d.add_paragraph()
    _p(d, "ATENȚIE: Această dispoziție de șantier nu se poate aplica fără semnătura proiectantului și a verificatorului Vg.", bold=True)
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_carte_tehnica() -> bytes:
    """Cartea Tehnică a Construcției — Lg 10/1995 + HG 273/1994."""
    d = Document()
    _h(d, "CARTEA TEHNICĂ A CONSTRUCȚIEI", size=14)
    _p(d, "Instalație gaze naturale — Lg 10/1995 + HG 273/1994", bold=True)
    d.add_paragraph()
    _p(d, "Beneficiar: {{beneficiar}}")
    _p(d, "Adresa lucrării: {{adresa_lucrare}}, {{localitate}}, jud. {{judet}}")
    _p(d, "Nr. proiect: {{numar_proiect}} / {{data_document}}")
    d.add_paragraph()
    _p(d, "CAPITOL A — DOCUMENTAȚIA PRIVIND PROIECTAREA", bold=True)
    _p(d, "A.1 Certificat urbanism nr. {{cu_numar}} / {{cu_data}}")
    _p(d, "A.2 Autorizație construire nr. {{ac_numar}} / {{ac_data}}")
    _p(d, "A.3 Proiect tehnic (PT) + Detalii execuție (DE)")
    _p(d, "A.4 Referat verificare tehnică Vg ({{verificator_vgd}} — {{atestat_vgd}})")
    _p(d, "A.5 Avize ANRE / OSD / Primărie / ROMGAZ / DELGAZ / ENEL (după caz)")
    _p(d, "A.6 Memoriu tehnic specialitate gaze")
    _p(d, "A.7 Caiet de sarcini și liste cantități")
    _p(d, "A.8 Devize generale")
    d.add_paragraph()
    _p(d, "CAPITOL B — DOCUMENTAȚIA PRIVIND EXECUȚIA", bold=True)
    _p(d, "B.1 Autorizație ANRE executant {{executant}} ({{atestat_executant}})")
    _p(d, "B.2 Decizia numire RTE {{responsabil_rte}} ({{autorizatie_rte}})")
    _p(d, "B.3 Procese verbale faze determinante:")
    _p(d, "    • PVRFD-01 Trasare traseu")
    _p(d, "    • PVRFD-02 Săpături + adâncime")
    _p(d, "    • PVRFD-03 Sudare/electrofuziune conductă PE/OL")
    _p(d, "    • PVRFD-04 Probe presiune (rezistență + etanșeitate)")
    _p(d, "    • PVRFD-05 Acoperire tranșee + bandă avertizare")
    _p(d, "B.4 Certificate calitate materiale și echipamente")
    _p(d, "B.5 Buletine analiză sudori atestați ANRE")
    _p(d, "B.6 Programul de control al calității execuției (PCCVI)")
    _p(d, "B.7 Jurnal de șantier")
    _p(d, "B.8 Dispoziții de șantier emise (lista anexată)")
    d.add_paragraph()
    _p(d, "CAPITOL C — DOCUMENTAȚIA PRIVIND RECEPȚIA", bold=True)
    _p(d, "C.1 PV recepție la terminarea lucrărilor nr. {{pvrt_numar}} / {{pvrt_data}}")
    _p(d, "C.2 PV punere în funcțiune emis de OSD {{osd}}")
    _p(d, "C.3 PV recepție finală nr. {{pvrf_numar}} / {{pvrf_data}}")
    _p(d, "C.4 Garanții și certificate de conformitate finale")
    d.add_paragraph()
    _p(d, "CAPITOL D — DOCUMENTAȚIA PRIVIND EXPLOATAREA, ÎNTREȚINEREA ȘI REVIZIILE", bold=True)
    _p(d, "D.1 Instrucțiuni de exploatare a instalației")
    _p(d, "D.2 Programul de revizii tehnice periodice (RTP)")
    _p(d, "D.3 Programul de verificare tehnică periodică (VTP)")
    _p(d, "D.4 Înregistrări revizii și verificări periodice")
    _p(d, "D.5 Evidența incidentelor și reparațiilor")
    d.add_paragraph()
    _p(d, "Cartea tehnică a fost completată de:", bold=True)
    _p(d, "Proiectant: {{proiectant}}")
    _p(d, "Executant: {{executant}}")
    _p(d, "RTE: {{responsabil_rte}}")
    _p(d, "Predată beneficiarului la data: {{data_predare_ct}}")
    _p(d, "Semnătură primire beneficiar: {{beneficiar}} ___________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_pv_receptie_terminare() -> bytes:
    """Proces verbal de recepție la terminarea lucrărilor (PVRT) — HG 343/2017."""
    d = Document()
    _h(d, "PROCES VERBAL DE RECEPȚIE LA TERMINAREA LUCRĂRILOR", size=14)
    _p(d, "nr. {{pvrt_numar}} din data {{pvrt_data}}", bold=True)
    d.add_paragraph()
    _p(d, "Conform HG 343/2017 — Regulamentul de recepție a lucrărilor de construcții și instalații aferente acestora.")
    d.add_paragraph()
    _p(d, "1. IDENTIFICAREA OBIECTIVULUI", bold=True)
    _p(d, "Denumire obiectiv: {{denumire_investitie}}")
    _p(d, "Adresa: {{adresa_lucrare}}, {{localitate}}, jud. {{judet}}")
    _p(d, "Autorizație construire: nr. {{ac_numar}} / {{ac_data}}")
    _p(d, "Beneficiar (investitor): {{beneficiar}}")
    _p(d, "Proiectant: {{proiectant}} ({{atestat_proiectant}})")
    _p(d, "Executant: {{executant}} ({{atestat_executant}})")
    _p(d, "RTE: {{responsabil_rte}} ({{autorizatie_rte}})")
    d.add_paragraph()
    _p(d, "2. COMISIA DE RECEPȚIE", bold=True)
    _p(d, "Președinte: {{presedinte_receptie}}")
    _p(d, "Membri: {{membri_receptie}}")
    _p(d, "Specialist domeniu instalații gaze: {{specialist_gaze}}")
    d.add_paragraph()
    _p(d, "3. DOCUMENTE EXAMINATE", bold=True)
    _p(d, "• Autorizația de construire și actele atașate")
    _p(d, "• Proiectul tehnic verificat de Vg")
    _p(d, "• Cartea tehnică a construcției (capitolele A + B)")
    _p(d, "• Buletinele de încercări (probe presiune, sudură)")
    _p(d, "• Certificatele de conformitate ale materialelor")
    _p(d, "• PV faze determinante semnate")
    d.add_paragraph()
    _p(d, "4. CONSTATĂRI", bold=True)
    _p(d, "{{constatari_receptie}}")
    d.add_paragraph()
    _p(d, "5. CONCLUZIA COMISIEI", bold=True)
    _p(d, "Comisia hotărăște: {{decizia_receptie}}")
    _p(d, "(ADMITE / RESPINGE / ADMITE CU OBSERVAȚII — se completează una din variante)")
    d.add_paragraph()
    _p(d, "6. SEMNĂTURI COMISIE", bold=True)
    _p(d, "Președinte: ____________________      Membri: ____________________ ____________________")
    _p(d, "Investitor: ____________________      Executant: ____________________")
    _p(d, "Proiectant: ____________________      RTE: ____________________      Specialist gaze: ____________________")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_pccvi_gaz() -> bytes:
    """Program de Control al Calității Execuției (PCCVI) — gaze naturale."""
    d = Document()
    _h(d, "PROGRAM DE CONTROL AL CALITĂȚII EXECUȚIEI (PCCVI)", size=14)
    _p(d, "Lucrări instalații gaze naturale", bold=True)
    d.add_paragraph()
    _p(d, "Lucrare: {{denumire_investitie}}")
    _p(d, "Beneficiar: {{beneficiar}}")
    _p(d, "Executant: {{executant}}")
    _p(d, "RTE: {{responsabil_rte}}")
    _p(d, "Verificator atestat Vg: {{verificator_vgd}}")
    d.add_paragraph()
    _p(d, "TABEL FAZE DETERMINANTE", bold=True)
    _p(d, "Nr. │ Faza │ Document de control │ Participanți │ Document emis")
    _p(d, "1 │ Trasare traseu │ Plan trasare + PV │ Beneficiar + Proiectant + RTE │ PVRFD-01")
    _p(d, "2 │ Săpături + adâncime conducta │ Fișă măsurători │ RTE + Inspector ISC │ PVRFD-02")
    _p(d, "3 │ Sudare/electrofuziune │ Buletin sudori + radiografii │ RTE + Sudor atestat ANRE │ PVRFD-03")
    _p(d, "4 │ Proba presiune rezistență │ Diagrama presiune 6h │ RTE + Reprezentant OSD │ PVRFD-04a")
    _p(d, "5 │ Proba etanșeitate │ Diagrama 24h │ RTE + Reprezentant OSD │ PVRFD-04b")
    _p(d, "6 │ Acoperire tranșee │ PV de control │ RTE + Beneficiar │ PVRFD-05")
    _p(d, "7 │ Probă tehnologică │ Buletin probă cu gaz │ RTE + OSD + ISCIR (după caz) │ PV PIF")
    d.add_paragraph()
    _p(d, "OBSERVAȚII:", bold=True)
    _p(d, "• Fiecare fază determinantă necesită oprirea lucrărilor până la semnarea PV.")
    _p(d, "• Verificatorul atestat ANRE Vg participă obligatoriu la fazele 3, 4 și 5.")
    _p(d, "• Programul este aprobat de proiectant și acceptat de RTE și beneficiar.")
    d.add_paragraph()
    _p(d, "Întocmit proiectant: {{proiectant}}            Acceptat RTE: {{responsabil_rte}}")
    _p(d, "Avizat verificator Vg: {{verificator_vgd}}    Aprobat beneficiar: {{beneficiar}}")
    _p(d, "Data: {{data_document}}")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def _build_caiet_sarcini_gaz() -> bytes:
    """Caiet de Sarcini — instalație gaze naturale."""
    d = Document()
    _h(d, "CAIET DE SARCINI — INSTALAȚIE GAZE NATURALE", size=14)
    d.add_paragraph()
    _p(d, "1. OBIECTUL CAIETULUI DE SARCINI", bold=True)
    _p(d, "Prezentul caiet de sarcini se aplică lucrării: {{tip_lucrare}} pentru beneficiarul {{beneficiar}}, situată la adresa {{adresa_lucrare}}, {{localitate}}, jud. {{judet}}.")
    d.add_paragraph()
    _p(d, "2. STANDARDE ȘI NORMATIVE APLICABILE", bold=True)
    _p(d, "• Lg 123/2012 — Legea energiei electrice și a gazelor naturale")
    _p(d, "• Ord. ANRE 89/2018 — Norma Tehnică Instalații Gaze")
    _p(d, "• NTPEE-2018 — Normativ tehnic proiectare execuție instalații gaze")
    _p(d, "• SR EN 12007 / SR EN 1555 / SR EN 13480 (conducte, fitting-uri)")
    _p(d, "• Lg 10/1995 — Calitatea în construcții")
    _p(d, "• P 118-99 — Norme protecție la foc")
    d.add_paragraph()
    _p(d, "3. MATERIALE ȘI ECHIPAMENTE", bold=True)
    _p(d, "• Conductă: {{material_conducta}}, dimensiune {{diametru_conducta}}, presiune nominală {{presiune_regim}}")
    _p(d, "• Robineți: cu sferă cu pasaj integral, certificați PE/OL")
    _p(d, "• Post reglare: {{post_reglare}}, cu valvă de siguranță SAV + SBV")
    _p(d, "• Contor: {{contor_orientativ}}, clasă metrologică G")
    _p(d, "• Toate materialele vor avea certificat de conformitate CE și declarație producător.")
    d.add_paragraph()
    _p(d, "4. CONDIȚII DE EXECUȚIE", bold=True)
    _p(d, "• Adâncime minimă montaj conductă: 0.9 m sub teren natural")
    _p(d, "• Bandă avertizare galbenă cu inscripție 'ATENȚIE GAZ' la 30 cm deasupra conductei")
    _p(d, "• Distanțe normate față de alte rețele: 0.4 m apă, 0.5 m electric, 1 m canalizare")
    _p(d, "• Sudarea va fi efectuată numai de sudori atestați ANRE cu autorizație în vigoare")
    d.add_paragraph()
    _p(d, "5. RECEPȚIA LUCRĂRILOR", bold=True)
    _p(d, "Recepția se face pe faze determinante conform PCCVI și se finalizează cu PV recepție la terminarea lucrărilor.")
    d.add_paragraph()
    _p(d, "Întocmit: {{proiectant}}                  Data: {{data_document}}")
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


SYSTEM_TEMPLATES = [
    {
        "key": "sys_cerere_racordare_gaz",
        "name": "Cerere racordare gaze naturale (sistem)",
        "industry": "gas_engineering",
        "subdomain": "bransamente_gaz",
        "builder": _build_cerere_racordare,
    },
    {
        "key": "sys_memoriu_tehnic_gaz",
        "name": "Memoriu tehnic instalație gaze (sistem)",
        "industry": "gas_engineering",
        "subdomain": "instalatii_utilizare",
        "builder": _build_memoriu_tehnic,
    },
    {
        "key": "sys_borderou_documente_gaz",
        "name": "Borderou documente dosar gaze (sistem)",
        "industry": "gas_engineering",
        "subdomain": "bransamente_gaz",
        "builder": _build_borderou,
    },
    {
        "key": "sys_adresa_osd_gaz",
        "name": "Adresă către OSD (sistem)",
        "industry": "gas_engineering",
        "subdomain": "bransamente_gaz",
        "builder": _build_adresa_osd,
    },
    {
        "key": "sys_certificare_vgd_gaz",
        "name": "Certificare internă VGD (sistem)",
        "industry": "gas_engineering",
        "subdomain": "bransamente_gaz",
        "builder": _build_certificare_vgd,
    },
    {
        "key": "sys_certificare_rte_gaz",
        "name": "Certificare internă RTE (sistem)",
        "industry": "gas_engineering",
        "subdomain": "bransamente_gaz",
        "builder": _build_certificare_rte,
    },
    # ============================================================
    # V5.7 — Faze complete documentare gaze naturale (DTAC/DTOE/CT/DS/PVRT/PCCVI/CS)
    # ============================================================
    {
        "key": "sys_dtac_gaz",
        "name": "D.T.A.C. — Documentație Autorizație Construire (sistem)",
        "industry": "gas_engineering",
        "subdomain": "autorizatii_construire",
        "builder": _build_dtac_gaz,
    },
    {
        "key": "sys_dtoe_gaz",
        "name": "D.T.O.E. — Organizare Execuție (sistem)",
        "industry": "gas_engineering",
        "subdomain": "executie_santier",
        "builder": _build_dtoe_gaz,
    },
    {
        "key": "sys_dispozitie_santier_gaz",
        "name": "Dispoziție de Șantier (sistem)",
        "industry": "gas_engineering",
        "subdomain": "executie_santier",
        "builder": _build_dispozitie_santier,
    },
    {
        "key": "sys_carte_tehnica_gaz",
        "name": "Carte Tehnică a Construcției (sistem)",
        "industry": "gas_engineering",
        "subdomain": "receptie_pif",
        "builder": _build_carte_tehnica,
    },
    {
        "key": "sys_pvrt_gaz",
        "name": "PV Recepție Terminare Lucrări (sistem)",
        "industry": "gas_engineering",
        "subdomain": "receptie_pif",
        "builder": _build_pv_receptie_terminare,
    },
    {
        "key": "sys_pccvi_gaz",
        "name": "PCCVI — Program Control Calitate Execuție (sistem)",
        "industry": "gas_engineering",
        "subdomain": "executie_santier",
        "builder": _build_pccvi_gaz,
    },
    {
        "key": "sys_caiet_sarcini_gaz",
        "name": "Caiet de Sarcini (sistem)",
        "industry": "gas_engineering",
        "subdomain": "instalatii_utilizare",
        "builder": _build_caiet_sarcini_gaz,
    },
    {
        "key": "sys_cerere_racordare_fv",
        "name": "Cerere racordare fotovoltaic (sistem)",
        "industry": "photovoltaics",
        "subdomain": "racordare_fv",
        "builder": _build_cerere_racordare_fv,
    },
    {
        "key": "sys_memoriu_tehnic_fv",
        "name": "Memoriu tehnic fotovoltaic — smart IF/ELSE (sistem)",
        "industry": "photovoltaics",
        "subdomain": "racordare_fv",
        "builder": _build_memoriu_tehnic_fv,
    },
    {
        "key": "sys_adresa_od_fv",
        "name": "Adresă către OD pentru racordare FV (sistem)",
        "industry": "photovoltaics",
        "subdomain": "racordare_fv",
        "builder": _build_adresa_od_fv,
    },
]


async def seed_system_templates(db):
    """Ensure all system templates exist (idempotent)."""
    now = datetime.now(timezone.utc).isoformat()
    from docx_processor import extract_placeholders
    for tpl in SYSTEM_TEMPLATES:
        existing = await db.system_templates.find_one({"key": tpl["key"]})
        data = tpl["builder"]()
        placeholders = extract_placeholders(data)
        doc = {
            "key": tpl["key"],
            "template_id": tpl["key"],
            "name": tpl["name"],
            "industry": tpl["industry"],
            "subdomain": tpl["subdomain"],
            "placeholders": placeholders,
            "size_bytes": len(data),
            "data_b64": base64.b64encode(data).decode("ascii"),
            "is_system": True,
            "created_at": existing["created_at"] if existing else now,
            "updated_at": now,
        }
        if existing:
            await db.system_templates.update_one({"key": tpl["key"]}, {"$set": doc})
        else:
            await db.system_templates.insert_one(doc)
